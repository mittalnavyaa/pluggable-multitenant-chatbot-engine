"""DOCX raw text extraction."""

from __future__ import annotations

from pathlib import Path
from zipfile import BadZipFile

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from exceptions.extraction_exceptions import CorruptedFileError
from extractors.base_extractor import BaseExtractor


class DOCXExtractor(BaseExtractor):
    """Extract raw text from Microsoft Word DOCX files."""

    supported_extension = ".docx"

    def _extract_text(self, file_path: Path) -> tuple[str, int]:
        try:
            document = Document(str(file_path))
            paragraph_text = [paragraph.text for paragraph in document.paragraphs]
            table_text = self._extract_tables(document)
            raw_text = "\n".join([*paragraph_text, *table_text]).strip()
            return raw_text, 1
        except (PackageNotFoundError, BadZipFile, KeyError) as exc:
            raise CorruptedFileError(f"Corrupted DOCX file: {file_path.name}") from exc
        except Exception as exc:
            raise CorruptedFileError(
                f"Unable to extract text from DOCX file: {file_path.name}"
            ) from exc

    @staticmethod
    def _extract_tables(document: Document) -> list[str]:
        rows: list[str] = []
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    rows.append("\t".join(cells))
        return rows
