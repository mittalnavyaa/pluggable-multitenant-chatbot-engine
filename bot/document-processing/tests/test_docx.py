"""Tests for DOCX text extraction."""

from __future__ import annotations

from pathlib import Path

from docx import Document

from extractors.docx_extractor import DOCXExtractor


def test_valid_docx_extracts_text(tmp_path: Path) -> None:
    file_path = tmp_path / "onboarding-guide.docx"
    document = Document()
    document.add_paragraph("Employee onboarding guide")
    document.save(str(file_path))

    result = DOCXExtractor().extract(file_path)

    assert result.success is True
    assert result.file_extension == ".docx"
    assert result.page_count == 1
    assert "Employee onboarding guide" in result.raw_text


def test_empty_docx_is_valid_with_zero_words(tmp_path: Path) -> None:
    file_path = tmp_path / "empty.docx"
    Document().save(str(file_path))

    result = DOCXExtractor().extract(file_path)

    assert result.success is True
    assert result.raw_text == ""
    assert result.word_count == 0


def test_corrupted_docx_returns_error(tmp_path: Path) -> None:
    file_path = tmp_path / "corrupted.docx"
    file_path.write_bytes(b"not a valid docx package")

    result = DOCXExtractor().extract(file_path)

    assert result.success is False
    assert result.errors
    assert "DOCX" in result.errors[0]
